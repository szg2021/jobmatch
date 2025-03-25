import os
import shutil
from typing import Any, Dict, List, Optional, Union, BinaryIO
from pathlib import Path
import logging

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import Document, Resume, Job, DocumentType, ProcessStatus
from app.schemas.document import DocumentCreate, DocumentUpdate, ResumeCreate, JobCreate
from app.services.recommendation_service import index_resume, index_job, update_resume_index, update_job_index

# 获取日志记录器
logger = logging.getLogger("app.services.document")

async def create_document(
    db: Session, 
    file: UploadFile, 
    user_id: int, 
    document_type: DocumentType
) -> Document:
    """
    创建文档记录并保存上传的文件
    """
    # 验证文件类型
    content_type = file.content_type
    if content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="仅支持PDF和DOCX格式")
    
    # 验证文件大小
    file_size = 0
    file.file.seek(0, 2)  # 移动到文件末尾
    file_size = file.file.tell()  # 获取文件大小
    file.file.seek(0)  # 重置文件指针
    
    if file_size > settings.MAX_DOCUMENT_SIZE * 1024 * 1024:  # 转换为字节
        raise HTTPException(status_code=400, detail=f"文件大小超过限制 ({settings.MAX_DOCUMENT_SIZE}MB)")
    
    # 创建上传目录（如果不存在）
    upload_dir = Path(settings.DOCUMENT_STORAGE_PATH)
    if not upload_dir.exists():
        upload_dir.mkdir(parents=True)
    
    # 生成文件路径
    filename = file.filename
    file_path = os.path.join(settings.DOCUMENT_STORAGE_PATH, f"{user_id}_{filename}")
    
    # 保存文件
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")
    
    # 创建文档记录
    db_document = Document(
        user_id=user_id,
        filename=filename,
        file_path=file_path,
        content_type=content_type,
        file_size=file_size,
        document_type=document_type,
        process_status=ProcessStatus.PENDING
    )
    
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    return db_document


async def process_document(db: Session, document_id: int) -> Document:
    """
    处理文档，提取文本并更新状态
    """
    # 获取文档
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 更新状态为处理中
    document.process_status = ProcessStatus.PROCESSING
    db.commit()
    
    try:
        # 提取文本
        text = await extract_text_from_document(document.file_path, document.content_type)
        
        # 更新文档记录
        document.original_text = text
        document.process_status = ProcessStatus.COMPLETED
        
        # 创建简历或职位记录
        if document.document_type == DocumentType.RESUME:
            # 使用外部AI解析简历信息（简化版实现）
            resume_data = {
                "document_id": document.id,
                "user_id": document.user_id,
                "title": "简历",  # 实际中应使用解析的标题
                "name": "",       # 实际中应使用解析的姓名
                "contact": "",    # 实际中应使用解析的联系方式
                "skills": "",     # 实际中应使用解析的技能
                "experience": "", # 实际中应使用解析的经验
                "education": ""   # 实际中应使用解析的教育背景
            }
            db_resume = Resume(**resume_data)
            db.add(db_resume)
            db.commit()
            db.refresh(db_resume)
            
            # 索引到向量数据库
            vector_id = await index_resume(db, db_resume, text)
            # 更新vector_id
            document.vector_id = vector_id
        
        elif document.document_type == DocumentType.JOB:
            # 使用外部AI解析职位信息（简化版实现）
            job_data = {
                "document_id": document.id,
                "user_id": document.user_id,
                "title": "职位",    # 实际中应使用解析的标题
                "company_name": "", # 实际中应使用解析的公司名称
                "location": "",     # 实际中应使用解析的地点
                "job_type": "",     # 实际中应使用解析的职位类型
                "salary_range": "", # 实际中应使用解析的薪资范围
                "requirements": "", # 实际中应使用解析的要求
                "responsibilities": "" # 实际中应使用解析的职责
            }
            db_job = Job(**job_data)
            db.add(db_job)
            db.commit()
            db.refresh(db_job)
            
            # 索引到向量数据库
            vector_id = await index_job(db, db_job, text)
            # 更新vector_id
            document.vector_id = vector_id
        
        db.commit()
        db.refresh(document)
        
    except Exception as e:
        # 处理失败
        document.process_status = ProcessStatus.FAILED
        db.commit()
        raise HTTPException(status_code=500, detail=f"处理文档失败: {str(e)}")
    
    return document


async def extract_text_from_document(file_path: str, content_type: str) -> str:
    """
    从文档中提取文本
    """
    try:
        if content_type == "application/pdf":
            # 处理PDF文件
            with fitz.open(file_path) as doc:
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
        
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # 处理DOCX文件
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n".join(paragraphs)
        
        else:
            raise ValueError(f"不支持的文件类型: {content_type}")
    
    except Exception as e:
        raise Exception(f"提取文本失败: {str(e)}")


async def get_documents_by_user(db: Session, user_id: int) -> List[Document]:
    """
    获取用户的所有文档
    """
    return db.query(Document).filter(Document.user_id == user_id).all()


async def get_document(db: Session, document_id: int) -> Optional[Document]:
    """
    获取单个文档
    """
    return db.query(Document).filter(Document.id == document_id).first()


async def get_resume_by_document(db: Session, document_id: int) -> Optional[Resume]:
    """
    通过文档ID获取简历
    """
    return db.query(Resume).filter(Resume.document_id == document_id).first()


async def get_job_by_document(db: Session, document_id: int) -> Optional[Job]:
    """
    通过文档ID获取职位
    """
    return db.query(Job).filter(Job.document_id == document_id).first()


async def get_resume(db: Session, resume_id: int) -> Optional[Resume]:
    """
    获取单个简历
    """
    return db.query(Resume).filter(Resume.id == resume_id).first()


async def get_job(db: Session, job_id: int) -> Optional[Job]:
    """
    获取单个职位
    """
    return db.query(Job).filter(Job.id == job_id).first()


async def get_resumes_by_user(db: Session, user_id: int) -> List[Resume]:
    """
    获取用户的所有简历
    """
    return db.query(Resume).filter(Resume.user_id == user_id).all()


async def get_jobs_by_user(db: Session, user_id: int) -> List[Job]:
    """
    获取用户的所有职位
    """
    return db.query(Job).filter(Job.user_id == user_id).all()


async def update_resume(db: Session, resume_id: int, update_data: Dict[str, Any]) -> Resume:
    """
    更新简历信息
    """
    resume = await get_resume(db, resume_id)
    if not resume:
        logger.warning(f"尝试更新不存在的简历: {resume_id}")
        raise HTTPException(status_code=404, detail="简历不存在")
    
    # 获取关联的文档
    document = await get_document(db, resume.document_id)
    if not document:
        logger.warning(f"简历{resume_id}关联的文档不存在: {resume.document_id}")
        raise HTTPException(status_code=404, detail="关联的文档不存在")
    
    # 更新简历信息
    for field, value in update_data.items():
        setattr(resume, field, value)
    
    db.commit()
    db.refresh(resume)
    
    # 如果文档已经有向量ID，则更新向量数据库
    if document.vector_id:
        try:
            logger.info(f"更新简历向量: resume_id={resume_id}, vector_id={document.vector_id}")
            await update_resume_index(
                vector_id=document.vector_id,
                resume=resume,
                document_text=document.original_text
            )
        except Exception as e:
            logger.error(f"更新简历向量失败: resume_id={resume_id}, error={str(e)}")
            # 这里我们选择不让向量更新失败影响API操作
            # 但在生产环境中可能需要更复杂的错误处理
    
    return resume


async def update_job(db: Session, job_id: int, update_data: Dict[str, Any]) -> Job:
    """
    更新职位信息
    """
    job = await get_job(db, job_id)
    if not job:
        logger.warning(f"尝试更新不存在的职位: {job_id}")
        raise HTTPException(status_code=404, detail="职位不存在")
    
    # 获取关联的文档
    document = await get_document(db, job.document_id)
    if not document:
        logger.warning(f"职位{job_id}关联的文档不存在: {job.document_id}")
        raise HTTPException(status_code=404, detail="关联的文档不存在")
    
    # 更新职位信息
    for field, value in update_data.items():
        setattr(job, field, value)
    
    db.commit()
    db.refresh(job)
    
    # 如果文档已经有向量ID，则更新向量数据库
    if document.vector_id:
        try:
            logger.info(f"更新职位向量: job_id={job_id}, vector_id={document.vector_id}")
            await update_job_index(
                vector_id=document.vector_id,
                job=job,
                document_text=document.original_text
            )
        except Exception as e:
            logger.error(f"更新职位向量失败: job_id={job_id}, error={str(e)}")
            # 这里我们选择不让向量更新失败影响API操作
            # 但在生产环境中可能需要更复杂的错误处理
    
    return job 